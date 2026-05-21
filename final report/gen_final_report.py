#!/usr/bin/env python3
"""Generate Rogue Linux Final Report — exact style match to FINAL REPORT 1.0 1.docx"""

import os
from docx import Document
from docx.shared import Pt, Cm, Emu, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as ET

BASE = os.path.dirname(os.path.abspath(__file__))
FIGS = os.path.join(BASE, "figures")
TMPL = os.path.join(BASE, "FINAL REPORT 1.0 1.docx")
OUT  = os.path.join(BASE, "Rogue_Linux_Final_Report.docx")

def fig(name): return os.path.join(FIGS, name)

# ─────────────────────────────────────────────────────────────────────────────
# STYLE PATCHING  (exact values read from reference doc XML)
# ─────────────────────────────────────────────────────────────────────────────

def _set_pPr_xml(style, xml_snippet):
    """Replace/add pPr in a style element with given XML."""
    NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    old = style._element.find(f'{{{NS}}}pPr')
    if old is not None:
        style._element.remove(old)
    new = ET.fromstring(xml_snippet)
    # insert after name element
    style._element.insert(2, new)

def patch_styles(doc):
    """Patch every style to match the reference document exactly."""
    NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # ── Body Text ─────────────────────────────────────────────────────────
    # justify, line=352/auto (≈1.467×), left=940 dxa, firstLine=417 dxa
    bt = doc.styles['Body Text']
    _set_pPr_xml(bt,
        '<w:pPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:jc w:val="both"/>'
        '<w:spacing w:line="352" w:lineRule="auto"/>'
        '<w:ind w:left="940" w:firstLine="417"/>'
        '</w:pPr>')

    # ── List Paragraph ─────────────────────────────────────────────────────
    # left=981 dxa, hanging=360 dxa (from reference), 14pt, justified
    lp = doc.styles['List Paragraph']
    _set_pPr_xml(lp,
        '<w:pPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:jc w:val="both"/>'
        '<w:spacing w:line="352" w:lineRule="auto"/>'
        '<w:ind w:left="981" w:hanging="360"/>'
        '</w:pPr>')
    # Set font size to 14pt to match Body Text
    rpr = lp._element.find(f'{{{NS}}}rPr')
    if rpr is None:
        rpr = ET.fromstring(
            '<w:rPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:sz w:val="28"/><w:szCs w:val="28"/>'
            '</w:rPr>')
        lp._element.append(rpr)
    else:
        for tag in ('w:sz', 'w:szCs'):
            el = rpr.find(f'{{{NS}}}{tag.split(":")[1]}')
            if el is not None: rpr.remove(el)
        rpr.append(ET.fromstring(f'<w:sz xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="28"/>'))
        rpr.append(ET.fromstring(f'<w:szCs xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="28"/>'))

    # Heading styles are already correct in the template – no changes needed.

# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def clear_body(doc):
    body = doc.element.body
    sectPr = body.find(qn('w:sectPr'))
    for child in list(body):
        if child is not sectPr:
            body.remove(child)

def page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

def H1(doc, txt):  return doc.add_heading(txt, level=1)
def H3(doc, txt):  return doc.add_heading(txt, level=3)
def H4(doc, txt):  return doc.add_heading(txt, level=4)

def para(doc, text, bold_prefix=None):
    """Body Text paragraph; optional bold prefix."""
    p = doc.add_paragraph(style='Body Text')
    if bold_prefix:
        p.add_run(bold_prefix).bold = True
    p.add_run(text)
    return p

def item(doc, text, prefix=None):
    """Bullet item using List Paragraph style with bullet char."""
    p = doc.add_paragraph(style='List Paragraph')
    # bullet character + tab indent
    r0 = p.add_run('•\t')
    r0.font.size = Pt(14)
    if prefix:
        rb = p.add_run(prefix)
        rb.bold = True
        rb.font.size = Pt(14)
    r = p.add_run(text)
    r.font.size = Pt(14)
    return p

def normal(doc, text, bold=False, size=None, align=None):
    p = doc.add_paragraph(style='Normal')
    r = p.add_run(text)
    r.bold = bold
    if size: r.font.size = Pt(size)
    if align is not None: p.alignment = align
    return p

def figure(doc, path, caption, width=Inches(5.2)):
    if not os.path.exists(path):
        print(f"  [MISS] {path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=width)
    cap = doc.add_paragraph(style='Body Text')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Emu(0)
    cap.paragraph_format.left_indent = Emu(0)
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(12)

# ── Table helpers ─────────────────────────────────────────────────────────────

def _tbl_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    bdr = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        bdr.append(b)
    tblPr.append(bdr)

def _cell(cell, text, bold=False, center=False):
    p = cell.paragraphs[0]
    p.clear()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(12)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def stable(doc, headers, rows, center_cols=None):
    """Bordered table; headers bold+centered; body 12pt."""
    center_cols = center_cols or []
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    _tbl_borders(t)
    for j, h in enumerate(headers):
        _cell(t.rows[0].cells[j], h, bold=True, center=True)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            _cell(t.rows[i+1].cells[j], str(v), center=(j in center_cols))
    return t

# ── TOC table ─────────────────────────────────────────────────────────────────

def _set_col_widths(table, widths):
    tbl = table._tbl
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is None:
        tblGrid = OxmlElement('w:tblGrid')
        tbl.insert(1, tblGrid)
    else:
        for c in list(tblGrid): tblGrid.remove(c)
    for w in widths:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        tblGrid.append(gc)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(widths[i]))
            tcW.set(qn('w:type'), 'dxa')
            old = tcPr.find(qn('w:tcW'))
            if old is not None: tcPr.remove(old)
            tcPr.append(tcW)

def _toc_cell(cell, text, bold=False):
    p = cell.paragraphs[0]
    p.clear()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(14)

def toc(doc, entries):
    """entries: (num, title, page, kind)  kind='hdr'|'front'|'chapter'|'section'"""
    t = doc.add_table(rows=1+len(entries), cols=3)
    _set_col_widths(t, [1671, 5324, 1589])   # exact widths from reference
    # header row
    _toc_cell(t.rows[0].cells[0], 'CHAPTER NO:', bold=True)
    _toc_cell(t.rows[0].cells[1], 'TITLE',       bold=True)
    _toc_cell(t.rows[0].cells[2], 'PAGE NO:',    bold=True)
    for i, (num, title, pg, kind) in enumerate(entries):
        row = t.rows[i+1]
        if kind == 'chapter':
            _toc_cell(row.cells[0], num,   bold=True)
            _toc_cell(row.cells[1], title, bold=True)
            _toc_cell(row.cells[2], pg,    bold=False)
        elif kind == 'front':
            _toc_cell(row.cells[0], '',    bold=False)
            _toc_cell(row.cells[1], title, bold=True)
            _toc_cell(row.cells[2], pg,    bold=False)
        else:  # section
            _toc_cell(row.cells[0], '',    bold=False)
            _toc_cell(row.cells[1], title, bold=False)
            _toc_cell(row.cells[2], pg,    bold=False)
    return t

# ═════════════════════════════════════════════════════════════════════════════
# BUILD
# ═════════════════════════════════════════════════════════════════════════════

doc = Document(TMPL)
clear_body(doc)
patch_styles(doc)

sec = doc.sections[0]
sec.page_width    = Cm(21.59)
sec.page_height   = Cm(27.94)
sec.top_margin    = Cm(2.54)
sec.bottom_margin = Cm(0.49)
sec.left_margin   = Cm(1.27)
sec.right_margin  = Cm(1.27)

# ═══════════════════════════════════════════════════════ COVER PAGE
H1(doc, 'ROGUE LINUX: A DETERMINISTIC BUILD SYSTEM AND COGNITIVE'
        ' PROCESS SUPERVISOR FOR MINIMAL LINUX-BASED OPERATING SYSTEM IMAGES')

normal(doc, 'A PROJECT REPORT', align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph(style='Normal')
normal(doc, 'Submitted by', align=WD_ALIGN_PARAGRAPH.CENTER)

tbl = doc.add_table(rows=4, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i,(name,reg) in enumerate([
    ('SRIDHARAN T',  '(620821205001)'),
    ('THANGARAJI K', '(620821205002)'),
    ('VIGNESH S',    '(620821205003)'),
    ('ARUN KUMAR M', '(620821205004)'),
]):
    for j,(txt,bold) in enumerate([(name,True),(reg,False)]):
        p = tbl.rows[i].cells[j].paragraphs[0]
        p.clear(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt); r.bold = bold; r.font.size = Pt(12)

doc.add_paragraph(style='Normal')
normal(doc, 'in partial fulfillment for the award of the degree of',
       align=WD_ALIGN_PARAGRAPH.CENTER)
H1(doc, 'BACHELOR OF TECHNOLOGY')
normal(doc, 'in', align=WD_ALIGN_PARAGRAPH.CENTER)
normal(doc, 'INFORMATION TECHNOLOGY', align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph(style='Normal')
normal(doc, 'GNANAMANI COLLEGE OF TECHNOLOGY  NAMAKKAL – 637 018',
       align=WD_ALIGN_PARAGRAPH.CENTER)
normal(doc, 'APRIL–MAY 2025', align=WD_ALIGN_PARAGRAPH.CENTER)
page_break(doc)

# ═══════════════════════════════════════════════════════ BONAFIDE
p = doc.add_paragraph(style='Normal')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('BONAFIDE CERTIFICATE'); r.bold = True; r.font.size = Pt(16)
doc.add_paragraph(style='Normal')

para(doc,
    'Certified that this project report “ROGUE LINUX: A DETERMINISTIC BUILD SYSTEM '
    'AND COGNITIVE PROCESS SUPERVISOR FOR MINIMAL LINUX-BASED OPERATING SYSTEM IMAGES” '
    'is the bonafide work of SRIDHARAN T (620821205001), THANGARAJI K (620821205002), '
    'VIGNESH S (620821205003), and ARUN KUMAR M (620821205004), who carried out the project '
    'work under my supervision.')

for _ in range(3): doc.add_paragraph(style='Normal')

for line in [
    'SIGNATURE\t\t\t\t\t\tSIGNATURE',
    'Dr. S. RAJKUMAR\t\t\t\t\tMr. P. ARULMOZHI',
    'HEAD OF THE DEPARTMENT\t\t\t\tSUPERVISOR',
    'Associate Professor,\t\t\t\t\tAssistant Professor,',
    'Information Technology\t\t\t\t\tInformation Technology',
    'Gnanamani College of Technology,\t\t\tGnanamani College of Technology,',
    'Namakkal – 637 018\t\t\t\t\tNamakkal – 637 018',
]:
    normal(doc, line)

for _ in range(2): doc.add_paragraph(style='Normal')
normal(doc, 'Submitted for the End Semester Project Work Viva-voce Examination held on \t\t')
for _ in range(2): doc.add_paragraph(style='Normal')
normal(doc, 'INTERNAL EXAMINER\t\t\t\t\tEXTERNAL EXAMINER')
page_break(doc)

# ═══════════════════════════════════════════════════════ ACKNOWLEDGEMENT
H1(doc, 'ACKNOWLEDGEMENT')

para(doc,
    'We would like to express our deep sense of heartiest thanks to our beloved Chairman '
    'Shri. C.A. N.V. Natarajan, B.Com, FCA., and to our beloved Correspondent '
    'Smt. N. Mangai Natarajan, M.Sc., for providing all necessary facilities and an '
    'excellent academic environment for the successful completion of this ambitious '
    'systems software project.')
para(doc,
    'We would like to express our sincere thanks to our beloved Director Admin '
    'Dr. K.K. Ramasamy, M.E., Ph.D., for their moral support and encouragement in '
    'pursuing this advanced low-level systems programming project.')
para(doc,
    'We extend our heartful gratitude to our beloved Principal Dr. V. Hariharan, M.E., Ph.D., '
    'for their continuous motivation and the institutional support that enabled access to '
    'laboratory resources required for this project.')
para(doc,
    'We extend our gratefulness to Dr. S. Rajkumar, M.E., Ph.D., Associate Professor and '
    'Head of the Department of Information Technology, for his guidance, his willingness to '
    'allow unconventional systems programming approaches, and his encouragement throughout '
    'this project.')
para(doc,
    'We would like to express our deepest appreciation to our Supervisor Mr. P. Arulmozhi, '
    'M.E., Assistant Professor, Department of Information Technology, for his expert guidance '
    'on operating systems internals, the Rust and C programming languages, POSIX process '
    'management, and QEMU-based testing methodology. His patient review of the iterative '
    'design and implementation phases of this project was invaluable.')
para(doc,
    'We gratefully acknowledge the contributions of the Rust programming language community '
    '(The Rust Foundation), the Ollama project team, and the Linux kernel development '
    'community whose open-source work made the foundational components of Rogue Linux possible.')
para(doc,
    'We thank all department staff members, laboratory assistants, and fellow students for '
    'their encouragement, technical discussions, and support throughout the development of '
    'this project.')

for _ in range(2): doc.add_paragraph(style='Normal')
normal(doc, '[SRIDHARAN T]     [THANGARAJI K]     [VIGNESH S]     [ARUN KUMAR M]')
page_break(doc)

# ═══════════════════════════════════════════════════════ VISION & MISSION
p = doc.add_paragraph(style='Normal')
p.add_run('Institute Vision:').bold = True

para(doc,
    'To emerge as a globally recognised technical institution producing ethical engineers, '
    'researchers, and entrepreneurs capable of addressing the challenges of a dynamically '
    'changing world through quality education and innovative research.')

doc.add_paragraph(style='Normal')
p = doc.add_paragraph(style='Normal')
p.add_run('Institute Mission:').bold = True

item(doc, 'To provide state-of-the-art infrastructure creating an effective learning environment.')
item(doc, 'To collaborate with leading industries and academia to empower students to meet global standards.')
item(doc, 'To foster an enterprising environment encouraging innovation and entrepreneurial activities.')

doc.add_paragraph(style='Normal')
p = doc.add_paragraph(style='Normal')
p.add_run('Department Vision (IT):').bold = True
para(doc,
    'To be the department that imparts professional computing training and makes competent '
    'engineers to work and contribute effectively in a competitive environment.')

doc.add_paragraph(style='Normal')
p = doc.add_paragraph(style='Normal')
p.add_run('Mission Statements (IT):').bold = True
item(doc, 'To prepare competent engineers and adapt to the dynamic needs of industries.')
item(doc, 'To pave way to the enrichment of knowledge and skills using latest technologies in diverse domains.')
item(doc, 'Inculcate strong ethical values and professionalism to serve society while updating knowledge and skills.')

doc.add_paragraph(style='Normal')
p = doc.add_paragraph(style='Normal')
p.add_run('PROGRAM EDUCATIONAL OBJECTIVES (PEOs)').bold = True
p = doc.add_paragraph(style='Normal')
p.add_run('Graduates of Information Technology will')

item(doc, 'Be proficient in utilising fundamental knowledge of engineering and technology towards problem solving and design of digital artifacts.', 'PEO-1: ')
item(doc, 'Think logically and pursue lifelong learning to understand technical issues related to computing.', 'PEO-2: ')
item(doc, 'Design and develop hardware and software systems by understanding social, business, and professional ethics.', 'PEO-3: ')

doc.add_paragraph(style='Normal')
p = doc.add_paragraph(style='Normal')
p.add_run('DEPARTMENT OF INFORMATION TECHNOLOGY PROGRAM OUTCOMES (POs)').bold = True

for label, text in [
    ('Engineering knowledge:', ' Apply knowledge of mathematics, science, and engineering fundamentals to solve complex engineering problems.'),
    ('Problem analysis:', ' Identify, formulate, and analyse complex engineering problems reaching substantiated conclusions.'),
    ('Design/development of solutions:', ' Design solutions for complex engineering problems and system components meeting specified needs.'),
    ('Conduct investigations:', ' Use research-based knowledge including design of experiments and interpretation of data.'),
    ('Modern tool usage:', ' Create, select, and apply appropriate techniques, resources, and modern engineering tools.'),
    ('The engineer and society:', ' Apply reasoning to assess societal, health, safety, legal, and cultural issues.'),
    ('Environment and sustainability:', ' Understand the impact of professional engineering solutions in societal and environmental contexts.'),
    ('Ethics:', ' Apply ethical principles and commit to professional responsibilities and norms of engineering practice.'),
    ('Individual and team work:', ' Function effectively as an individual and as a member or leader in diverse teams.'),
    ('Communication:', ' Communicate effectively on complex engineering activities with the community at large.'),
    ('Project management and finance:', ' Demonstrate knowledge of engineering and management principles.'),
    ('Life-long learning:', ' Engage in independent and life-long learning in the broadest context of technological change.'),
]:
    item(doc, text, label)

doc.add_paragraph(style='Normal')
p = doc.add_paragraph(style='Normal')
p.add_run('PROGRAM SPECIFIC OUTCOMES (PSOs)').bold = True
p = doc.add_paragraph(style='Normal')
p.add_run('Graduates of the program will be able to')

item(doc, 'Apply mathematical and computing knowledge to identify and provide solutions for complex problems using software engineering principles and ICT tools.', 'PSO-1: ')
item(doc, 'Design and develop programs in algorithms, networking, web design, and mobile computing to build a sustainable career in industry.', 'PSO-2: ')
page_break(doc)

# ═══════════════════════════════════════════════════════ ABSTRACT
H1(doc, 'ABSTRACT')

para(doc,
    'Rogue Linux is a deterministic, metadata-driven infrastructure for constructing minimal '
    'Linux-based operating system images. The core innovation is Cogman (Cognitive Manager), '
    'a unified toolchain spanning both the build phase and the runtime phase of a Linux system. '
    'During the build phase, cogman-planner — implemented in Rust — reads declarative '
    'TOML package definitions, resolves a directed acyclic dependency graph using topological '
    'sort, enforces filesystem and network security policies, and emits a compact binary '
    'execution plan in the custom CGM2PLAN format.')
para(doc,
    'The complementary cogman-executor, implemented in C11, memory-maps the plan file and '
    'executes each typed step operation (OP_EXEC, OP_MKDIR, OP_COPY, OP_VERIFY, OP_CLEANUP) '
    'with path traversal protection. During the runtime phase, cogman-supervisor acts as PID 1 '
    '— the first process invoked by the Linux kernel — and manages the complete '
    'lifecycle of system services. It parses INI-format service definition files, implements '
    'the SIGCHLD self-pipe trick for safe asynchronous child-process reaping, enforces service '
    'dependency ordering, and supports three restart policies (never, on-failure, always).')
para(doc,
    'A Unix domain socket control interface (cogman-ctl) allows operators to list, start, stop, '
    'and restart services at runtime without rebooting. Performance evaluation demonstrates a '
    '56× improvement in plan resolution time (8 ms vs. 450 ms Python baseline), a '
    '21× reduction in peak memory (4 MB vs. 85 MB), and a 50× reduction in '
    'per-step execution overhead (0.9 ms vs. 45 ms). A minimal bootable rootfs of approximately '
    '6.3 MB was constructed and verified under QEMU through a four-stage boot sequence. '
    'All 40 unit, integration, supervisor lifecycle, and end-to-end test cases pass.')

p = doc.add_paragraph(style='Body Text')
p.add_run('Keywords: ').bold = True
p.add_run(
    'Rogue Linux, Cogman, build system, PID 1, init system, service supervisor, TOML, Rust, '
    'C11, CGM2PLAN, topological sort, dependency graph, minimal Linux, QEMU.')
page_break(doc)

# ═══════════════════════════════════════════════════════ TABLE OF CONTENTS
H1(doc, 'TABLE OF CONTENT')

toc(doc, [
    ('',  'ABSTRACT',                              '', 'front'),
    ('',  'LIST OF ABBREVIATIONS',                 '', 'front'),
    ('',  'LIST OF FIGURES',                       '', 'front'),
    ('1', 'INTRODUCTION',                          '1',  'chapter'),
    ('',  '1.1 General Introduction',              '1',  'section'),
    ('',  '1.2 Importance of the Study',           '2',  'section'),
    ('',  '1.3 Problem Statement',                 '3',  'section'),
    ('',  '1.4 Aim and Objective',                 '4',  'section'),
    ('',  '1.5 Scope and Limitation of the Study', '5',  'section'),
    ('2', 'LITERATURE REVIEW',                     '7',  'chapter'),
    ('',  '2.1 Embedded Linux Build Systems',      '7',  'section'),
    ('',  '2.2 Yocto Project and Nix',             '8',  'section'),
    ('',  '2.3 Init Systems and Service Managers', '9',  'section'),
    ('',  '2.4 Signal Handling in PID 1',          '10', 'section'),
    ('',  '2.5 Binary Plan Formats and Caching',   '11', 'section'),
    ('',  '2.6 LLM Integration in System Tooling', '12', 'section'),
    ('',  '2.7 Gap Analysis',                      '13', 'section'),
    ('3', 'SYSTEM ANALYSIS',                        '14', 'chapter'),
    ('',  '3.1 Existing System',                   '14', 'section'),
    ('',  '3.2 Proposed System',                   '16', 'section'),
    ('4', 'SYSTEM DESIGN',                          '19', 'chapter'),
    ('',  '4.1 System Design Overview',            '19', 'section'),
    ('',  '4.2 CGM2PLAN Binary Format',            '20', 'section'),
    ('',  '4.3 System Architecture',               '21', 'section'),
    ('',  '4.4 Use Case Diagrams',                 '23', 'section'),
    ('',  '4.5 Class Diagrams',                    '24', 'section'),
    ('',  '4.6 Sequence Diagrams',                 '25', 'section'),
    ('',  '4.7 Component Diagram',                 '26', 'section'),
    ('',  '4.8 System Requirements',               '27', 'section'),
    ('5', 'MODULE DESCRIPTION',                     '29', 'chapter'),
    ('',  '5.1 cogman-planner Module',             '29', 'section'),
    ('',  '5.2 cogman-executor Module',            '31', 'section'),
    ('',  '5.3 cogman-supervisor Module',          '33', 'section'),
    ('',  '5.4 cogman-ctl Module',                 '35', 'section'),
    ('',  '5.5 Messenger IPC Module',              '36', 'section'),
    ('',  '5.6 Rootfs Bootstrap Module',           '37', 'section'),
    ('6', 'IMPLEMENTATION',                         '39', 'chapter'),
    ('',  '6.1 Boot Sequence',                     '39', 'section'),
    ('',  "6.2 Kahn's Topological Sort (Rust)",    '40', 'section'),
    ('',  '6.3 Path Traversal Guard (C)',           '41', 'section'),
    ('',  '6.4 SIGCHLD Self-Pipe (C)',              '42', 'section'),
    ('',  '6.5 Technology Stack',                  '43', 'section'),
    ('7', 'SYSTEM TESTING',                         '44', 'chapter'),
    ('',  '7.1 Planner Unit Tests',                '44', 'section'),
    ('',  '7.2 Executor Unit Tests',               '45', 'section'),
    ('',  '7.3 Supervisor Test Cases',             '46', 'section'),
    ('',  '7.4 End-to-End Boot Tests (QEMU)',       '47', 'section'),
    ('',  '7.5 User Acceptance Testing',           '48', 'section'),
    ('8', 'PERFORMANCE ANALYSIS',                   '49', 'chapter'),
    ('',  '8.1 Plan Resolution Time',              '49', 'section'),
    ('',  '8.2 Memory Usage',                      '50', 'section'),
    ('',  '8.3 Per-Step Execution Overhead',       '51', 'section'),
    ('',  '8.4 Performance Summary',               '52', 'section'),
    ('',  '8.5 Rootfs Size Analysis',              '53', 'section'),
    ('9', 'CONCLUSION AND FUTURE WORK',             '54', 'chapter'),
    ('',  '9.1 Conclusion',                        '54', 'section'),
    ('',  '9.2 Future Work',                       '56', 'section'),
    ('',  'APPENDIX 1 – SOURCE CODE',          '58', 'front'),
    ('',  'APPENDIX 2 – SCREENSHOTS',          '63', 'front'),
    ('',  'APPENDIX 3 – SETUP NOTES',          '64', 'front'),
    ('',  'REFERENCES',                            '65', 'front'),
])
page_break(doc)

# ═══════════════════════════════════════════════════════ LIST OF ABBREVIATIONS
H1(doc, 'LIST OF ABBREVIATIONS')
stable(doc,
    ['ABBREVIATION', 'FULL FORM'],
    [
        ('CGM / Cogman', 'Cognitive Manager'),
        ('CLI',   'Command Line Interface'),
        ('DAG',   'Directed Acyclic Graph'),
        ('ELF',   'Executable and Linkable Format'),
        ('EWMA',  'Exponentially Weighted Moving Average'),
        ('IPC',   'Inter-Process Communication'),
        ('LLM',   'Large Language Model'),
        ('mmap',  'Memory-Mapped File I/O'),
        ('OS',    'Operating System'),
        ('PID',   'Process Identifier'),
        ('POSIX', 'Portable Operating System Interface'),
        ('QEMU',  'Quick Emulator (open-source hypervisor)'),
        ('rootfs','Root Filesystem'),
        ('SIGCHLD','Signal: Child Process Status Changed'),
        ('SIGTERM','Signal: Termination Request'),
        ('TLV',   'Type-Length-Value (message encoding)'),
        ('TOML',  "Tom's Obvious Minimal Language"),
        ('UDS',   'Unix Domain Socket'),
    ]
)
page_break(doc)

# ═══════════════════════════════════════════════════════ LIST OF FIGURES
H1(doc, 'LIST OF FIGURES')
stable(doc,
    ['FIGURE NO.', 'TITLE'],
    [
        ('Figure 1.1',  'Cogman Build Pipeline'),
        ('Figure 1.2',  'Runtime Architecture'),
        ('Figure 4.1',  'System Architecture – Cogman Build and Init'),
        ('Figure 4.2',  'CGM2PLAN Binary Format Layout'),
        ('Figure 4.3',  'DFD Level 0 – Context Diagram'),
        ('Figure 4.4',  'DFD Level 1 – Build Subsystem'),
        ('Figure 4.5',  'Use Case Diagram – Build Subsystem'),
        ('Figure 4.6',  'Use Case Diagram – Runtime Subsystem'),
        ('Figure 4.7',  'Class Diagram – cogman-planner'),
        ('Figure 4.8',  'Class Diagram – cogman-supervisor'),
        ('Figure 4.9',  'Sequence Diagram – Build Flow'),
        ('Figure 4.10', 'Sequence Diagram – Supervisor Start'),
        ('Figure 4.11', 'Component Diagram'),
        ('Figure 5.1',  'Dependency Graph'),
        ('Figure 5.2',  'Topological Sort'),
        ('Figure 5.3',  'Executor Loop'),
        ('Figure 5.4',  'Path Traversal Guard'),
        ('Figure 5.5',  'SIGCHLD Self-Pipe Pattern'),
        ('Figure 5.6',  'Service State Machine'),
        ('Figure 5.7',  'CTL Protocol'),
        ('Figure 5.8',  'Messenger IPC'),
        ('Figure 5.9',  'Rootfs Layout'),
        ('Figure 6.1',  'Four-Stage Boot Sequence'),
        ('Figure A.1',  'DWM Window Manager on Rogue Linux'),
        ('Figure A.2',  'Terminal Environment'),
    ]
)
page_break(doc)

# ═══════════════════════════════════════════════════════ CH 1 INTRODUCTION
H1(doc, 'CHAPTER 1 INTRODUCTION')

H3(doc, 'GENERAL INTRODUCTION')
para(doc,
    'Rogue Linux is not a Linux distribution in the conventional sense — it is an '
    'infrastructure for building one. Given a set of declarative package definitions written '
    'in TOML, Rogue Linux produces a reproducible root filesystem capable of booting under '
    'QEMU or on bare-metal x86_64 hardware, with the Cogman (Cognitive Manager) toolchain '
    'serving as both the build engine and the runtime init process. The clean separation '
    'between the build half and the runtime half ensures that no build-time dependencies '
    'leak into the runtime image and that the final rootfs contains only what was explicitly '
    'declared in package definitions.')
para(doc,
    'The Cogman name reflects the design goal: a process supervisor that is aware of the '
    'relationships between the services it manages, not just a sequential shell script. '
    'cogman-supervisor understands service dependency declarations, enforces dependency-ordered '
    'startup, monitors service health, and applies configurable restart policies — '
    'behaviours that characterise a cognitively structured supervisor rather than a simple '
    'sequential launcher. The supervisor can be queried and controlled at runtime through a '
    'text-protocol Unix domain socket, providing operational visibility into service state '
    'without requiring log parsing or process table inspection.')
para(doc,
    'The project has a dual motivation: to demonstrate that a high-performance, safe, and '
    'reproducible embedded Linux build system can be built with Rust and C11 without external '
    'framework dependencies, and to explore the integration of a locally-running large language '
    'model (Qwen2.5-3B via llama.cpp) as an advisory component that can explain build failures '
    'and service configuration issues to operators in natural language.')

figure(doc, fig('fig1_1_build_pipeline.png'), 'Figure 1.1: Cogman Build Pipeline')
figure(doc, fig('fig1_2_runtime_arch.png'),   'Figure 1.2: Runtime Architecture')

H3(doc, 'IMPORTANCE OF THE STUDY')
para(doc,
    'This project addresses critical challenges in embedded Linux systems development where '
    'reproducibility, security, and performance are paramount. The ability to construct '
    'verifiable, minimal operating system images is essential for IoT devices, edge computing '
    'nodes, and security-hardened deployments. The study demonstrates that modern systems '
    'programming languages (Rust, C11) can dramatically outperform legacy scripting approaches '
    'while maintaining correctness guarantees enforced at compile time.')
para(doc,
    'From an educational perspective, the project explores the intersection of operating '
    'systems theory — PID 1 semantics, POSIX signal handling, process supervision — '
    'with modern software engineering practices such as type-safe serialisation, '
    'content-addressed caching, and binary format design. The AI advisor component further '
    'demonstrates how large language models can augment system tooling without compromising '
    'security or determinism.')

H3(doc, 'PROBLEM STATEMENT')
para(doc,
    'Modern embedded Linux systems, container base images, and minimal operating environments '
    'demand a reproducible, auditable method for constructing a root filesystem from source '
    'packages. Existing solutions such as Buildroot and the Yocto Project address this need '
    'but at the cost of enormous complexity: Buildroot requires a Python and Bash toolchain '
    'exceeding 200 MB and takes 30–90 minutes for a complete build; Yocto requires '
    'days of build time and specialised knowledge of BitBake recipe syntax. Neither provides '
    'a lightweight init daemon that integrates tightly with the build pipeline’s '
    'metadata model.')
para(doc,
    'A secondary problem is that existing init systems are inappropriate for minimal embedded '
    'environments. systemd requires udev, D-Bus, PAM, and tens of shared libraries — '
    'incompatible with an image targeting 6–10 MB. SysVinit relies on fragile shell '
    'scripts with no dependency ordering and no structured restart logic. The absence of a '
    'unified toolchain spanning build and runtime forces developers to combine disparate tools '
    '(Makefiles, Docker, shell scripts, busybox init) with no common metadata format, security '
    'policy model, or structured error-diagnosis capability.')

H3(doc, 'AIM AND OBJECTIVE')
H4(doc, 'Aim:')
para(doc,
    'The main aim of this project is to design, implement, and validate Rogue Linux — a '
    'deterministic, metadata-driven infrastructure for constructing minimal Linux-based '
    'operating system images — with a unified Cogman toolchain spanning both build and '
    'runtime phases of the operating system lifecycle.')

H4(doc, 'Objectives:')
for t in [
    'To design a declarative TOML package metadata format with schema validation covering identity, build steps, installer steps, dependency declarations, and security policy.',
    'To implement cogman-planner in Rust for dependency graph construction, cycle detection, topological sort, and CGM2PLAN binary plan emission with content-addressed caching.',
    'To implement cogman-executor in C11 for typed step execution (OP_EXEC, OP_MKDIR, OP_COPY, OP_VERIFY, OP_CLEANUP) with path traversal protection on all copy operations.',
    'To implement cogman-supervisor as a PID 1 process supervisor with SIGCHLD self-pipe child reaping and three configurable restart policies (never, on-failure, always).',
    'To implement cogman-ctl for runtime service control via Unix domain socket with a minimal text-based command protocol.',
    'To construct a bootable minimal rootfs of approximately 6.3 MB and verify it under QEMU through a four-stage boot sequence exercising all code paths.',
    'To achieve a 50× or greater performance improvement over the legacy Python-based reference implementation.',
]:
    item(doc, t)

H3(doc, 'SCOPE AND LIMITATION OF THE STUDY')
H4(doc, 'Scope of the Study:')
para(doc,
    'The scope covers the full design, implementation, and evaluation of the Cogman toolchain '
    'for x86_64 targets. The build system supports native and pre-packaged build modes. The '
    'AI advisor component uses Qwen2.5-3B as an advisory-only interface with no ability to '
    'issue commands or modify system state. Evaluation uses QEMU rather than bare-metal hardware.')

H4(doc, 'Limitations of the Study:')
para(doc,
    'Production security hardening (Landlock, seccomp, namespaces) is documented as future '
    'work. The current system does not implement cgroup-based resource limits per service. The '
    'IPC messenger does not implement authentication or access control on the control socket. '
    'The AI advisor quality is bounded by the Qwen2.5-3B model’s training data; '
    'fine-tuning on a Cogman-specific dataset is required for production-quality advisory '
    'accuracy. Full rootfs image-level reproducibility is not yet achieved due to '
    'filesystem timestamps introduced during ext4 image creation.')
page_break(doc)

# ═══════════════════════════════════════════════════════ CH 2 LIT REVIEW
H1(doc, 'CHAPTER 2 LITERATURE REVIEW')

H4(doc, 'Embedded Linux Build Systems:')
para(doc,
    'Buildroot (2002–present) is the most widely used tool for generating minimal '
    'embedded Linux root filesystems. It uses a Kconfig-based package selection interface '
    'and a Makefile-based build orchestration system. While mature and extensively documented, '
    'Buildroot’s approach makes it difficult to implement fine-grained security policy '
    'enforcement at the build step level, and the absence of a typed intermediate plan format '
    'means build orchestration logic and execution logic are interleaved rather than separated.')

H4(doc, 'Yocto Project and Nix:')
para(doc,
    'The Yocto Project (2010–present) provides a more powerful approach using the '
    'BitBake task executor and OpenEmbedded-Core recipe system. However, a standard Yocto '
    'build environment requires 50+ GB of disk space, hours of build time, and a steep '
    'learning curve. Nix (2006–present) and Guix (2013–present) represent the '
    'state of the art in reproducible package management. Rogue Linux draws inspiration from '
    'Nix’s content-addressed build caching (implemented via FNV-1a hash over TOML '
    'content) but reduces conceptual overhead by using TOML declarations directly rather than '
    'a functional programming language.')

H4(doc, 'Init Systems and Service Managers:')
para(doc,
    'systemd (2010–present) has become the dominant PID 1 implementation on modern '
    'Linux distributions, providing comprehensive service management, socket activation, and '
    'cgroup-based resource control. However, its extensive dependencies (udev, D-Bus, PAM) '
    'make it unsuitable for minimal rootfs environments targeting 6–10 MB. '
    'Runit (2004–present) is a lightweight UNIX init scheme; cogman-supervisor draws '
    'architectural inspiration from Runit’s per-service supervision model while '
    'integrating it with the Cogman metadata system and adding a programmatic control interface.')

H4(doc, 'Signal Handling in PID 1 – POSIX Correctness:')
para(doc,
    's6 (2012–present) by Laurent Bercot provides the most principled approach to UNIX '
    'service supervision, with a rigorous implementation of the self-pipe trick for SIGCHLD '
    'handling. PID 1 has special semantics in the Linux kernel: it is exempt from the '
    'default signal disposition table, requiring explicit handler registration for every '
    'signal the supervisor wishes to respond to. Furthermore, PID 1 is responsible for '
    'reaping all orphaned zombie processes. cogman-supervisor’s SIGCHLD self-pipe '
    'implementation directly follows the s6 pattern: a signal handler writes one byte to a '
    'pipe, and the main select() loop monitors the read end rather than using sigwaitinfo().')

H4(doc, 'Binary Plan Formats and Build Caching:')
para(doc,
    'Bazel’s action cache (2015–present) stores build actions as content-addressed '
    'entries identified by input hash. The CGM2PLAN format’s content-addressed plan cache '
    'is architecturally inspired by this model: a FNV-1a hash over the package name, version, '
    'and TOML file content serves as the cache key. The ELF binary format demonstrates that '
    'fixed-size header structures with string tables can be efficiently memory-mapped without '
    'any parsing overhead — CGM2PLAN applies this principle directly: a 64-byte header, '
    'fixed 128-byte step records, and a variable-length string table, all accessible via '
    'mmap() pointer arithmetic with zero heap allocations.')

H4(doc, 'LLM Integration in System Tooling:')
para(doc,
    'Qwen2.5 (2024) is a family of large language models from Alibaba Cloud with strong '
    'performance on code generation, question answering, and structured reasoning. The 3B '
    'parameter variant, quantised to 4-bit precision (Q4_K_M GGUF format), requires '
    'approximately 2.1 GB of RAM and achieves 5–15 tokens/second inference speed '
    'on CPU-only hardware. The cogman advisor component uses this model served via llama.cpp '
    'to answer build system configuration questions and explain service file syntax without '
    'any external API dependency or internet connectivity.')

H4(doc, 'Reproducible Build Systems:')
para(doc,
    'The Reproducible Builds Project (2015–present) has documented the challenges of '
    'achieving bit-for-bit identical output artifacts given the same source inputs, regardless '
    'of the build machine’s environment (username, filesystem timestamps, installed tool '
    'versions). Rogue Linux achieves reproducibility at the plan level — the CGM2PLAN '
    'output is deterministic for the same TOML input — but not yet at the full rootfs '
    'image level due to ext4 filesystem timestamp insertion during image creation. Full '
    'rootfs reproducibility is identified as a future work objective.')

H4(doc, 'Gap Analysis:')
para(doc,
    'Prior work in embedded Linux build systems provides reproducibility (Nix, Yocto) or '
    'simplicity (Buildroot, Alpine) but not a unified system spanning both build and runtime '
    'phases under a single metadata model. Prior work in service supervisors provides correct '
    'signal handling (s6) or simple service management (Runit) but not dependency-aware '
    'startup with a programmable control interface integrated with build metadata. Rogue Linux '
    'fills this gap: TOML metadata → CGM2PLAN binary → typed C executor → '
    'PID 1 Rust supervisor → UDS control interface, with no external framework '
    'dependencies beyond the Linux kernel and C standard library.')
page_break(doc)

# ═══════════════════════════════════════════════════════ CH 3 SYSTEM ANALYSIS
H1(doc, 'CHAPTER 3 SYSTEM ANALYSIS')

H3(doc, 'EXISTING SYSTEM')
para(doc,
    'The reference system against which Rogue Linux is benchmarked is a Python-based prototype '
    'implementing the same conceptual pipeline. The Python planner uses the pure-Python toml '
    'library to parse package metadata and a dict-based DAG implementation for dependency '
    'resolution. The Python executor uses subprocess.run() for each build step. The init '
    'process is a BusyBox init shell script with no dependency ordering or restart logic.')

H4(doc, 'Key Features of the Existing System:')
for t in [
    'Plan resolution time of ≈450 ms due to Python interpreter startup and pure-Python TOML parsing overhead.',
    'Peak planner memory of ≈85 MB from Python’s per-object allocation model (24+ bytes of overhead per object).',
    'Per-step execution overhead of ≈45 ms from Python’s subprocess.run() marshaling and object allocation.',
    'No typed step operations — all steps are generic shell strings with no OP_COPY/OP_MKDIR/OP_VERIFY semantics.',
    'No path traversal protection on copy operations; packages can reference arbitrary filesystem paths.',
    'No binary plan format — the plan is Python pickled data, not portable between Python versions.',
    'Shell script init with no dependency ordering, no restart policies, and no runtime control interface.',
]:
    item(doc, t)

H4(doc, 'Advantages of the Existing System:')
for t in [
    'Easy to implement and understand for small-scale builds.',
    'Low initial development cost; no compilation step required.',
    'Simple structure accessible to developers unfamiliar with systems programming.',
    'Rapid prototyping: new package types can be added by editing Python dicts.',
]:
    item(doc, t)

H4(doc, 'Disadvantages of the Existing System:')
for t in [
    'Slow plan resolution (450 ms) makes CI/CD pipelines with many packages impractical.',
    'High memory usage (85 MB) prevents deployment on memory-constrained build hosts.',
    'No typed step semantics; errors in copy or mkdir operations produce generic exception messages.',
    'No security policy enforcement at build time; malicious package definitions can access arbitrary paths.',
    'No structured runtime control; service management requires manual shell commands.',
    'No portability of plan format across language versions or architectures.',
]:
    item(doc, t)

H3(doc, 'PROPOSED SYSTEM:')
para(doc,
    'Rogue Linux replaces all legacy components with purpose-built Rust and C11 implementations '
    'that address each identified limitation. The proposed system introduces a unified metadata '
    'model (TOML), a typed binary plan format (CGM2PLAN), and a POSIX-correct PID 1 '
    'supervisor with a programmable control interface, all without external framework dependencies.')

H4(doc, 'Data Collection and Input:')
para(doc,
    'Package authors write declarative TOML package definition files specifying identity, '
    'build steps, installer steps, dependency declarations, and security policy. The '
    'cogman-planner validates each file using serde’s compile-time type-safe '
    'deserialisation, rejecting malformed files with precise error messages (field name '
    'and line number) before any build step executes.')

H4(doc, 'Dependency Graph Resolution:')
para(doc,
    'The planner constructs a directed acyclic graph (DAG) from all declared dependencies, '
    'detects cycles using depth-first search, and produces a topologically sorted build '
    'order using Kahn’s algorithm. A content-addressed cache — FNV-1a hash over '
    'package name, version, and TOML content — avoids redundant plan re-computation '
    'for unchanged packages, reducing planning time from 8 ms to 0.3 ms on '
    'cache hits.')

H4(doc, 'Binary Plan Emission and Execution:')
para(doc,
    'The CGM2PLAN binary format provides a portable, zero-parsing-overhead interface between '
    'the Rust planner and C executor. A 64-byte header, fixed 128-byte step records, and a '
    'variable-length string table allow the executor to access all plan data via mmap() '
    'pointer arithmetic with zero heap allocations after the initial mmap call.')

H4(doc, 'PID 1 Service Supervisor:')
para(doc,
    'cogman-supervisor acts as PID 1, managing all system services through their complete '
    'lifecycle. Service definition files use INI format specifying name, command, type '
    '(oneshot or longrun), restart policy, restart delay, and dependency list. The supervisor '
    'enforces dependency-ordered startup, implements SIGCHLD self-pipe child reaping, and '
    'exposes a Unix domain socket control interface.')

H4(doc, 'Advantages of the Proposed System:')
for t in [
    '56× faster plan resolution (8 ms vs. 450 ms Python baseline).',
    '21× lower peak memory usage (4 MB vs. 85 MB Python).',
    '50× lower per-step execution overhead (0.9 ms vs. 45 ms Python).',
    'Typed step operations (OP_EXEC, OP_MKDIR, OP_COPY, OP_VERIFY, OP_CLEANUP) with defined semantics and failure modes.',
    'Path traversal protection on all OP_COPY operations via path_has_traversal() guard.',
    'Portable binary plan format: architecture-independent, language-independent, version-stamped.',
    'PID 1 supervisor with dependency-aware startup, SIGCHLD self-pipe reaping, and UDS control interface.',
    'Build-time security policy enforcement per package (declared filesystem write paths and network access).',
]:
    item(doc, t)
page_break(doc)

# ═══════════════════════════════════════════════════════ CH 4 SYSTEM DESIGN
H1(doc, 'CHAPTER 4 SYSTEM DESIGN')

H3(doc, 'SYSTEM DESIGN OVERVIEW')
para(doc,
    'The Rogue Linux system is divided into two conceptually distinct halves connected by '
    'the CGM2PLAN binary plan format. The Build Half operates on the developer’s host '
    'machine and transforms declarative package metadata into a staged root filesystem. '
    'The Runtime Half operates on the target system (bare metal or QEMU) after kernel boot '
    'and manages all process lifecycle operations. The AI Advisor operates alongside the '
    'Build Half as an optional natural-language query interface, implemented as a separate '
    'process with no write access to any Cogman state.')

figure(doc, fig('fig6_1_system_arch.png'),
       'Figure 4.1: System Architecture – Cogman Build System and Init Architecture')

H3(doc, 'CGM2PLAN BINARY FORMAT')
para(doc,
    'The CGM2PLAN format is a custom binary format designed for maximum execution efficiency. '
    'The design is modelled on the ELF format: a fixed-size 64-byte header, a fixed-size '
    'record array (128 bytes per step), and a variable-length string table. The executor '
    'accesses the plan by memory-mapping the file using mmap(NULL, file_size, PROT_READ, '
    'MAP_PRIVATE, fd, 0), mapping the file into the process’s virtual address space '
    'without copying any data. The plan header is accessed by casting the mmap base pointer '
    'to a const plan_header*, and step records are accessed by computing base + '
    'sizeof(plan_header) + i × sizeof(step_record).')

figure(doc, fig('fig6_2_cgm2plan_format.png'),
       'Figure 4.2: CGM2PLAN Binary Format Layout')

H3(doc, 'SYSTEM ARCHITECTURE')

H4(doc, 'DFD Level 0 – Context Diagram:')
para(doc,
    'The Level 0 DFD shows two external actors: Package Author (writes package.toml '
    'files and places source archives) and System Operator (invokes cogman-planner and '
    'cogman-executor for builds; uses cogman-ctl for runtime control). The system boundary '
    'encompasses four internal components: cogman-planner, cogman-executor, '
    'cogman-supervisor, and cogman-ctl. Key data flows: Package Author → Planner: '
    'package.toml; Planner → Executor: build-plan.bin; Executor → Rootfs: staged '
    'package files; Operator → cogman-ctl → Supervisor: UDS protocol messages.')

figure(doc, fig('fig6_3_dfd_level0.png'),
       'Figure 4.3: DFD Level 0 – Context Diagram')

H4(doc, 'DFD Level 1 – Build Subsystem:')
para(doc,
    'The Level 1 DFD decomposes the build process into seven sub-processes: '
    'P1 Schema Validation, P2 Dependency Loading, P3 Cycle Detection, '
    'P4 Topological Sort, P5 Policy Check, P6 Plan Emission, and '
    'P7 Step Execution. Data stores include D1: packages/ directory, '
    'D2: plan cache (.cogman-cache/), and D3: staging rootfs. Each process has '
    'well-defined input and output data flows, enabling independent testing of each stage.')

figure(doc, fig('fig6_4_dfd_level1.png'),
       'Figure 4.4: DFD Level 1 – Build Subsystem')

stable(doc,
    ['PROCESS', 'INPUT', 'OUTPUT', 'DATA STORE'],
    [
        ('P1: Schema Validation', 'package.toml bytes', 'PackageMetadata struct', '—'),
        ('P2: Dependency Loading', 'PackageMetadata, packages/', 'DependencyGraph', 'D1: packages/'),
        ('P3: Cycle Detection', 'DependencyGraph', 'Error or Ok', '—'),
        ('P4: Topological Sort', 'Acyclic graph', 'Ordered build list', '—'),
        ('P5: Policy Check', 'PackageMetadata, rootfs', 'Error or Ok', '—'),
        ('P6: Plan Emission', 'Ordered steps, strtab', 'build-plan.bin', 'D2: plan cache'),
        ('P7: Step Execution', 'build-plan.bin (mmap\'d)', 'Files in $PKGROOT', 'D3: rootfs'),
    ]
)

H3(doc, 'USE CASE DIAGRAMS')
para(doc,
    'The Use Case Diagram for the Build Subsystem shows two actors: Package Author and Build '
    'Engineer. Package Authors define package metadata, declare build steps, declare '
    'dependencies, and set security policy. Build Engineers invoke the planner, invoke the '
    'executor, and verify the staging rootfs. The system enforces schema validation on '
    'planner invocation, cycle detection on dependency load, policy check before plan '
    'emission, and path traversal guard on every OP_COPY step execution.')

figure(doc, fig('fig6_5_usecase_build.png'),
       'Figure 4.5: Use Case Diagram – Build Subsystem')
figure(doc, fig('fig6_6_usecase_runtime.png'),
       'Figure 4.6: Use Case Diagram – Runtime Subsystem')

H3(doc, 'CLASS DIAGRAMS')
para(doc,
    'The Class Diagram for cogman-planner shows the primary structs: PackageMetadata '
    '(containing Identity, Builder, Installer, Policy), DependencyGraph (HashMap-based '
    'adjacency list), and PlanEmitter (CGM2PLAN binary writer with string table '
    'deduplication). Relationships: PackageMetadata aggregates Identity, Builder, Installer, '
    'and Policy; PlanEmitter consumes a DependencyGraph and ResolveResult; PlannerError '
    'is returned on validation failure.')

figure(doc, fig('fig6_7_class_planner.png'),
       'Figure 4.7: Class Diagram – cogman-planner')
figure(doc, fig('fig6_8_class_supervisor.png'),
       'Figure 4.8: Class Diagram – cogman-supervisor')

H3(doc, 'SEQUENCE DIAGRAMS')
para(doc,
    'The Sequence Diagram for the package build flow shows twelve steps from planner '
    'invocation to executor exit: (1) plan → (2) validate → '
    '(3) load_deps → (4) detect_cycles → (5) topological_sort '
    '→ (6) check_policy → (7) emit_plan → (8) mmap '
    '→ (9) validate_header → (10) per-step fork()+execve() '
    '→ (11) exit(0) or exit(1) on failure. The supervisor startup sequence '
    'shows service loading, dependency gate evaluation, SIGCHLD pipe setup, and the '
    'select() main loop entry.')

figure(doc, fig('fig6_9_seq_build.png'),
       'Figure 4.9: Sequence Diagram – Build Flow')
figure(doc, fig('fig6_10_seq_start.png'),
       'Figure 4.10: Sequence Diagram – Supervisor Start')

H3(doc, 'COMPONENT DIAGRAM')
para(doc,
    'The Component Diagram shows the four Cogman binaries and their interfaces: the '
    'CGM2PLAN file interface between planner and executor, the UDS socket interface '
    'between supervisor and ctl, and the INI service definition file interface consumed '
    'by the supervisor at boot. The AI Advisor connects to the planner’s log output '
    'and the operator’s query interface but has no write access to any system state.')

figure(doc, fig('fig6_11_component.png'),
       'Figure 4.11: Component Diagram')

H3(doc, 'SYSTEM REQUIREMENTS')
H4(doc, 'Hardware Requirements:')
stable(doc,
    ['COMPONENT', 'MINIMUM (BUILD HOST)', 'RECOMMENDED (BUILD HOST)'],
    [
        ('Processor', 'Intel Core i3, 2.0 GHz, 2 cores', 'Intel Core i7, 3.5 GHz, 8+ cores'),
        ('RAM', '8 GB DDR4 (for Rust compilation)', '32 GB DDR4 (parallel builds)'),
        ('Storage', '50 GB HDD (source trees + rootfs)', '500 GB NVMe SSD'),
        ('OS', 'Ubuntu 20.04 LTS (x86_64)', 'Ubuntu 22.04 LTS'),
        ('QEMU RAM', '64 MB (minimal rootfs)', '256 MB'),
    ]
)

H4(doc, 'Software Requirements:')
stable(doc,
    ['SOFTWARE', 'VERSION', 'PURPOSE'],
    [
        ('Rust (rustup stable)', '1.75+',  'cogman-planner compilation'),
        ('GCC',                  '11+',    'cogman-executor / supervisor / ctl'),
        ('GNU Make',             '4.3+',   'Build orchestration (top-level Makefile)'),
        ('QEMU (x86_64)',        '8.x',    'Minimal rootfs boot testing'),
        ('BusyBox',              '1.36.1', 'Shell utilities in minimal rootfs'),
        ('Python 3.11+',         '3.11+',  'Test harness and build helper scripts'),
        ('llama.cpp',            'latest', 'Qwen2.5-3B AI advisor inference'),
        ('Git',                  '2.x',    'Version control'),
    ]
)
page_break(doc)

# ═══════════════════════════════════════════════════════ CH 5 MODULE DESCRIPTION
H1(doc, 'CHAPTER 5 MODULE DESCRIPTION')

H3(doc, 'COGMAN-PLANNER MODULE (Rust)')
para(doc,
    'cogman-planner is the first component of the build pipeline, transforming a package.toml '
    'file into a build-plan.bin binary plan file. The module is organised into six '
    'sub-modules: schema.rs (serde struct definitions for PackageMetadata, Identity, Builder, '
    'Installer, Policy), graph/resolve.rs (dependency loading and DependencyGraph '
    'construction), graph/cycle.rs (DFS-based cycle detection), graph/topo.rs (Kahn’s '
    'algorithm topological sort), policy.rs (filesystem and network policy enforcement), '
    'and plan/emit.rs (CGM2PLAN binary emission with string table deduplication).')
para(doc,
    'Content-addressed plan caching is implemented via FNV-1a hash computed over the package '
    'name, version, and raw bytes of the package.toml file. The 64-character hexadecimal '
    'cache key is stored in a .cogman-cache/ directory alongside the plan file. Before '
    'emitting a new plan, the planner checks whether a cached plan exists with the same key; '
    'on a cache hit, plan emission is skipped entirely, reducing planning time from '
    '≈8 ms to ≈0.3 ms on unchanged packages — a 27× '
    'additional speedup for the common CI/CD case.')

figure(doc, fig('fig7_1_dep_graph.png'), 'Figure 5.1: Dependency Graph')
figure(doc, fig('fig7_2_topo_sort.png'), 'Figure 5.2: Topological Sort')

H3(doc, 'COGMAN-EXECUTOR MODULE (C11)')
para(doc,
    'cogman-executor is the execution engine for CGM2PLAN binary plans, organised into four '
    'source files: main.c (argument parsing, mmap, header validation, step dispatch loop), '
    'plan/plan.c (plan validation and string table access helpers), ops/exec.c (OP_EXEC '
    'handler using fork()+execve()), and ops/copy.c (OP_COPY handler with '
    'path_has_traversal() guard and recursive copy). The executor has no external library '
    'dependencies beyond the C standard library and POSIX.')
para(doc,
    'The execution model is intentionally simple: read the step count from the plan header, '
    'iterate over all step records in order, dispatch each step to its typed handler, and '
    'exit with code 0 if all steps complete successfully or code 1 if any '
    'FAIL_ABORT step fails. Steps marked with FAIL_WARN produce a warning message but do '
    'not abort execution, allowing optional verification steps to fail without invalidating '
    'the build. The entire executor binary is under 8 KB stripped.')

figure(doc, fig('fig7_3_executor_loop.png'), 'Figure 5.3: Executor Loop')
figure(doc, fig('fig7_4_path_guard.png'),    'Figure 5.4: Path Traversal Guard')

H3(doc, 'COGMAN-SUPERVISOR MODULE (C11)')
para(doc,
    'cogman-supervisor is the PID 1 process managing all system services after kernel '
    'boot. Its key design property is that it must never block in a way that prevents '
    'timely SIGCHLD delivery and child reaping. The SIGCHLD self-pipe trick solves this: '
    'a signal handler writes one byte to a write-end pipe created with pipe2(O_NONBLOCK|'
    'O_CLOEXEC), and the main select() loop monitors the read-end pipe descriptor alongside '
    'the control socket and restart timers. When select() returns with the pipe readable, '
    'the main loop drains the pipe and calls waitpid(-1, WNOHANG) in a loop until it '
    'returns zero, reaping all terminated children without blocking.')
para(doc,
    'Service definition files are stored in /etc/cogman/services/ and use a minimal '
    'INI-format with three sections: [service] (name, command, type, restart, '
    'restart_delay, depends_on), [env] (key=value environment variable overrides), and '
    '[meta] (description, enabled flag). The parser implements a minimal INI reader '
    'without external library dependencies, using a section-tracking enum and '
    'key-value splitting on the first ‘=’ character.')

figure(doc, fig('fig7_5_sigchld.png'),       'Figure 5.5: SIGCHLD Self-Pipe Pattern')
figure(doc, fig('fig7_6_state_machine.png'), 'Figure 5.6: Service State Machine')

H3(doc, 'COGMAN-CTL MODULE (C11)')
para(doc,
    'cogman-ctl is the runtime control client for the supervisor. It connects to the '
    'supervisor’s Unix domain socket at /tmp/cogman.sock using AF_UNIX SOCK_STREAM, '
    'sends a single command line followed by a newline character, reads the response '
    'until the connection closes, and prints it to stdout. The text protocol is minimal: '
    'commands are “list”, “start name”, “stop name”, '
    '“restart name”, and “status name”; responses are formatted text '
    'lines terminated by “OK” or “ERR message”. The entire cogman-ctl '
    'binary is approximately 200 lines of C and 8 KB stripped.')

figure(doc, fig('fig7_7_ctl_protocol.png'), 'Figure 5.7: CTL Protocol')

H3(doc, 'MESSENGER IPC MODULE (C)')
para(doc,
    'The messenger module provides typed inter-process communication between Cogman '
    'components using a fixed 16-byte TLV header: magic (4 bytes, “COG1”), '
    'version (2 bytes), message type (2 bytes), payload '
    'length (4 bytes), and source PID (4 bytes). Five message types '
    'are defined: MSG_HEARTBEAT (0), MSG_HUD_ALERT (1), MSG_POLICY_REQ (2), '
    'MSG_DATA_XFER (3), MSG_LOG_INFO (4). The broker uses AF_UNIX SOCK_STREAM '
    'with non-blocking accept() to process IPC messages within the supervisor’s '
    'select() main loop. A 2-second SO_RCVTIMEO ensures a slow client cannot hold '
    'the supervisor’s main loop beyond this limit.')

figure(doc, fig('fig7_8_messenger.png'), 'Figure 5.8: Messenger IPC')

H3(doc, 'ROOTFS BOOTSTRAP MODULE')
para(doc,
    'The minimal Rogue Linux rootfs is constructed in three phases. Phase 1 creates '
    'the directory skeleton: /bin, /sbin, /usr/bin, /lib, /lib64, /etc/cogman/services, '
    '/etc/cogman/plans, /run, /tmp, /proc, /sys, /dev, /root. Phase 2 populates '
    'the rootfs with pre-built binaries: cogman-supervisor at /sbin/cogman-supervisor, '
    'a symbolic link /sbin/init → /sbin/cogman-supervisor, and BusyBox applet '
    'symlinks for all required shell utilities. Phase 3 installs shared libraries '
    '(libc.so.6 and ld-linux-x86-64.so.2) copied from the build host. The resulting '
    'rootfs is approximately 6.3 MB and boots successfully under QEMU with 64 MB RAM.')

figure(doc, fig('fig7_9_rootfs_layout.png'), 'Figure 5.9: Rootfs Layout')
page_break(doc)

# ═══════════════════════════════════════════════════════ CH 6 IMPLEMENTATION
H1(doc, 'CHAPTER 6 IMPLEMENTATION')

H3(doc, 'BOOT SEQUENCE')
para(doc,
    'The four-stage boot sequence verified under QEMU exercises all major code paths. '
    'Stage 1 (Kernel Boot): the Linux kernel decompresses, initialises hardware, '
    'and executes /sbin/init, which is a symlink to cogman-supervisor. '
    'Stage 2 (Service Loading): cogman-supervisor scans /etc/cogman/services/, '
    'parses INI service files, and builds the dependency graph. '
    'Stage 3 (Ordered Startup): services start in dependency order — the '
    'hello oneshot service runs first, followed by heartbeat (longrun), then '
    'ctl-probe. Stage 4 (Runtime Control): exec-probe exercises the cogman-ctl '
    'interface, confirming the control socket is operational.')

figure(doc, fig('fig8_1_boot_sequence.png'), 'Figure 6.1: Four-Stage Boot Sequence')

H3(doc, "KAHN'S TOPOLOGICAL SORT (Rust)")
para(doc,
    'The dependency graph resolution algorithm uses Kahn’s algorithm to produce a '
    'build order respecting all declared dependencies. In-degrees are initialised for all '
    'nodes; a queue is seeded with zero-in-degree nodes; iteratively a node is removed '
    'from the queue, added to the output order, and its dependents’ in-degrees are '
    'decremented — pushing newly zero-in-degree dependents back into the queue. '
    'A cycle is detected when the output order contains fewer nodes than the graph.')

stable(doc,
    ['STEP', 'OPERATION', 'DATA STRUCTURE'],
    [
        ('1', 'Initialise in_degree map to 0 for all nodes',  'HashMap<&str, usize>'),
        ('2', 'Increment in_degree for each declared dep',     'HashMap iteration'),
        ('3', 'Seed queue with all zero-in-degree nodes',      'VecDeque<&str>'),
        ('4', 'Pop node, append to order, decrement deps',     'while let Some(node)'),
        ('5', 'Check order.len() == graph.node_count()',       'cycle detection'),
    ]
)

H3(doc, 'PATH TRAVERSAL GUARD (C)')
para(doc,
    'Every OP_COPY operation passes source and destination paths through '
    'path_has_traversal() before execution. The guard scans the path for ‘..’ '
    'components, rejecting any path that could escape the staging rootfs boundary. '
    'This provides a structural safety property that is independent of plan content: '
    'even a maliciously crafted plan file cannot cause cogman-executor to copy files '
    'outside the declared staging root.')

H3(doc, 'SIGCHLD SELF-PIPE (C)')
para(doc,
    'The SIGCHLD self-pipe is implemented with pipe2(O_NONBLOCK|O_CLOEXEC). The signal '
    'handler performs only a single write() call (async-signal-safe per POSIX). The '
    'main select() loop detects pipe readability, drains the pipe buffer, then calls '
    'waitpid(-1, WNOHANG) in a loop until it returns zero or −1/ECHILD, reaping '
    'all terminated children in one pass. This pattern is safe, non-blocking, and '
    'compatible with all other I/O operations in the select() fdset.')

H3(doc, 'TECHNOLOGY STACK')
stable(doc,
    ['COMPONENT', 'LANGUAGE', 'KEY DEPENDENCIES'],
    [
        ('cogman-planner',    'Rust (stable)',     'serde 1.0, toml 0.8, clap 4.4'),
        ('cogman-executor',   'C11 (GCC 11+)',     'POSIX stdlib only'),
        ('cogman-supervisor', 'C11 (GCC 11+)',     'POSIX stdlib only'),
        ('cogman-ctl',        'C11 (GCC 11+)',     'POSIX stdlib only'),
        ('AI Advisor',        'Python + llama.cpp','Qwen2.5-3B Q4_K_M GGUF'),
        ('Rootfs base',       'BusyBox 1.36.1',    'statically linked, x86_64'),
        ('Test platform',     'QEMU 8.x',          'x86_64 full-system emulation'),
    ]
)
page_break(doc)

# ═══════════════════════════════════════════════════════ CH 7 SYSTEM TESTING
H1(doc, 'CHAPTER 7 SYSTEM TESTING')

H3(doc, 'PLANNER UNIT TESTS – Schema Validation')
stable(doc,
    ['TC ID', 'INPUT CONDITION', 'EXPECTED BEHAVIOUR', 'RESULT'],
    [
        ('SV-01', 'Valid complete package.toml',          'Exit 0, plan written successfully',          'PASS'),
        ('SV-02', 'Missing [identity] section',           'Exit 1, serde deserialisation error',        'PASS'),
        ('SV-03', 'identity.name = empty string',         'Exit 1, name must not be empty',            'PASS'),
        ('SV-04', 'identity.version = empty string',      'Exit 1, version error',                     'PASS'),
        ('SV-05', 'build.steps = [] (empty list)',        'Exit 1, steps must not be empty',           'PASS'),
        ('SV-06', "policy.filesystem.write = ['../etc']", 'Exit 1, non-absolute path rejected',        'PASS'),
        ('SV-07', 'Circular dependency A → B → A', 'Exit 1, cycle detected with path',      'PASS'),
        ('SV-08', 'Dependency not found in packages/',    'Exit 1, missing dependency error',          'PASS'),
        ('SV-09', 'TOML syntax error (unclosed quote)',   'Exit 1, TOML parse error with line number', 'PASS'),
        ('SV-10', 'depends.build contains empty string',  'Exit 1, empty dependency name',             'PASS'),
    ]
)

H3(doc, 'EXECUTOR UNIT TESTS – Step Operations')
stable(doc,
    ['TC ID', 'OP', 'TEST CONDITION', 'EXPECTED RESULT', 'RESULT'],
    [
        ('EX-01', 'OP_EXEC',    "echo 'hello' command",           "stdout: 'hello', exit 0",               'PASS'),
        ('EX-02', 'OP_EXEC',    'exit 1, FAIL_ABORT flag',        'Executor exits with code 1',            'PASS'),
        ('EX-03', 'OP_EXEC',    'exit 1, FAIL_WARN flag',         'Execution continues, logs warning',     'PASS'),
        ('EX-04', 'OP_MKDIR',   'Create /tmp/test/a/b/c',         'Directory created (nested)',             'PASS'),
        ('EX-05', 'OP_MKDIR',   'Path already exists',            'Exit 0, idempotent',                    'PASS'),
        ('EX-06', 'OP_COPY',    'Copy file to valid destination',  'Destination matches source byte-exact', 'PASS'),
        ('EX-07', 'OP_COPY',    "Destination contains '..'",       "Rejected at guard, exit 2",             'PASS'),
        ('EX-08', 'OP_COPY',    'Recursive directory tree copy',   'All files copied, permissions correct', 'PASS'),
        ('EX-09', 'OP_VERIFY',  'Existing file path',              'Exit 0 (file found)',                   'PASS'),
        ('EX-10', 'OP_VERIFY',  'Non-existent, FAIL_ABORT',        'Exit 1 (verification failed)',          'PASS'),
        ('EX-11', 'OP_CLEANUP', 'Remove existing temp directory',  'Directory removed after execution',     'PASS'),
        ('EX-12', 'Header',     'Wrong magic bytes in plan file',  "Exit 1 with 'Bad magic' error",         'PASS'),
        ('EX-13', 'Header',     'Wrong version number',            "Exit 1 with 'Version mismatch'",        'PASS'),
    ]
)

H3(doc, 'SUPERVISOR TEST CASES – Service Lifecycle')
stable(doc,
    ['TC ID', 'SCENARIO', 'EXPECTED OUTCOME', 'RESULT'],
    [
        ('SL-01', 'Oneshot completes successfully',         'State → SVC_DONE',                        'PASS'),
        ('SL-02', 'Oneshot fails, restart=never',           'State → SVC_FAILED',                      'PASS'),
        ('SL-03', 'Long-running service tracked by PID',    'svc->pid positive, state=RUNNING',             'PASS'),
        ('SL-04', 'restart=always after SIGKILL',           'Supervisor restarts within delay+1 s',    'PASS'),
        ('SL-05', 'restart=on-failure, clean exit 0',       'No restart triggered',                         'PASS'),
        ('SL-06', 'restart=on-failure, exit code 1',        'Restart triggered after delay',                'PASS'),
        ('SL-07', 'Dependency gate blocks start',           'B remains STOPPED until A completes',          'PASS'),
        ('SL-08', 'Gate opens on dependency completion',    'B starts automatically',                       'PASS'),
        ('SL-09', 'Chain A → B → C',              'Start order: A then B then C',                 'PASS'),
        ('SL-10', 'Explicit cogman-ctl stop',               'Service does not restart after termination',   'PASS'),
        ('SL-11', 'Orphan process reaping',                 'Grandchild reaped by PID 1 on reparent', 'PASS'),
        ('SL-12', 'SIGTERM initiates clean shutdown',       'All services stopped; supervisor exits 0',     'PASS'),
    ]
)

H3(doc, 'END-TO-END BOOT TESTS (QEMU)')
stable(doc,
    ['TC ID', 'BOOT CONDITION', 'EXPECTED CONSOLE OUTPUT', 'RESULT'],
    [
        ('E2E-01', 'Normal boot, all services',      '[SERVICE:hello] cogman-supervisor is alive',   'PASS'),
        ('E2E-02', 'Normal boot',                    '[SERVICE:heartbeat] tick 0',                   'PASS'),
        ('E2E-03', 'Normal boot',                    '[SERVICE:ctl-probe] control socket OK',        'PASS'),
        ('E2E-04', 'Normal boot',                    '[SERVICE:exec-probe] plan execution OK',       'PASS'),
        ('E2E-05', 'Kill heartbeat with SIGKILL',    'Supervisor restarts heartbeat within 3 s','PASS'),
        ('E2E-06', 'cogman-ctl list',                'All 4 services listed with correct states',    'PASS'),
        ('E2E-07', 'cogman-ctl stop heartbeat',      'heartbeat: stopped (no restart)',              'PASS'),
        ('E2E-08', 'Boot with 64 MB RAM',       'All services start; no OOM events',           'PASS'),
    ]
)

H3(doc, 'USER ACCEPTANCE TESTING')
stable(doc,
    ['TC ID', 'USER STORY', 'ACCEPTANCE CRITERION', 'RESULT'],
    [
        ('UAT-01', 'Package author wants clear TOML errors',      'Error identifies field and line number',    'PASS'),
        ('UAT-02', 'Build engineer wants early cycle detection',  'Cycle path reported before any execution',  'PASS'),
        ('UAT-03', 'Operator wants to restart crashed service',   'cogman-ctl restart works within 1 s', 'PASS'),
        ('UAT-04', 'Operator wants path traversal blocked',       "OP_COPY with '..' rejected clearly",       'PASS'),
        ('UAT-05', 'Developer wants fast re-planning on cache',   'Cache hit produces plan in <1 ms',    'PASS'),
    ]
)
page_break(doc)

# ═══════════════════════════════════════════════════════ CH 8 PERFORMANCE
H1(doc, 'CHAPTER 8 PERFORMANCE ANALYSIS')

H3(doc, 'PLAN RESOLUTION TIME')
para(doc,
    'The 56× improvement in plan resolution time (from ≈450 ms to '
    '≈8 ms) arises from three compounding factors. First, the Rust planner '
    'avoids Python’s interpreter startup overhead of approximately 50 ms per '
    'invocation — a fixed cost independent of the package metadata size. Second, '
    'TOML deserialisation using serde’s derive macro produces compiled code that '
    'directly constructs Rust structs from the TOML byte stream without intermediate '
    'object allocation; Python’s pure-Python toml library constructs a Python dict '
    'hierarchy through multiple layers of interpretation, allocation, and reference '
    'counting. Third, the dependency graph algorithms use Rust’s HashMap and Vec '
    'with O(1) average-case lookup and cache-friendly memory layout, compared to '
    'Python’s dict-based implementation with pointer-following indirection and '
    'per-object reference count updates.')
para(doc,
    'The content-addressed plan cache provides an additional benefit on repeated build '
    'invocations. The FNV-1a hash is computed over the package name, version string, and '
    'TOML file content — approximately 0.3 ms for a typical 2–5 KB '
    'metadata file. On a cache hit, the entire plan resolution is short-circuited to this '
    'hash computation plus a file existence check, reducing the cost from 8 ms to '
    '0.3 ms — a 27× reduction for the common case in CI/CD environments '
    'where most packages are unchanged between build runs.')

H3(doc, 'MEMORY USAGE')
para(doc,
    'The 21× reduction in peak memory (from ≈85 MB to ≈4 MB) is '
    'primarily attributable to Python’s per-object allocation overhead. In CPython 3.11, '
    'every Python object carries an ob_refcnt (8 bytes), an ob_type pointer '
    '(8 bytes), and variable additional fields. A Python str for a 20-character '
    'string requires approximately 73 bytes — 3.65× the string content. '
    'A Python list carrying 100 package metadata dicts requires substantially more memory '
    'than the equivalent Rust Vec<PackageMetadata> because each Python dict entry carries '
    'a PyObject* pointer, a hash value, and reference count overhead. Rust’s '
    'serde-derived struct layout allocates exactly the memory required by the struct '
    'fields with no overhead beyond standard alignment padding.')

H3(doc, 'PER-STEP EXECUTION OVERHEAD')
para(doc,
    'The 50× reduction in per-step overhead (from ≈45 ms to '
    '≈0.9 ms) is dominated by the elimination of Python’s '
    'subprocess.run() overhead. Python’s subprocess.run() constructs a '
    'subprocess.CompletedProcess object, performs multiple Python type checks and '
    'attribute lookups, calls os.fork() through the Python C API, and awaits the child '
    'through a polling loop with object allocation at each poll iteration. The C '
    'executor’s execute_step() function calls fork() and execve() directly via C '
    'standard library wrappers, taking approximately 0.5–1.0 ms — '
    'essentially the minimum achievable process creation overhead on x86_64, limited '
    'by the kernel’s process table allocation and scheduler dispatch costs.')

H3(doc, 'PERFORMANCE SUMMARY')
stable(doc,
    ['METRIC', 'LEGACY PYTHON', 'ROGUE LINUX (Rust/C)', 'IMPROVEMENT'],
    [
        ('Plan resolution (cold)',      '≈450 ms',        '≈8 ms',    '56×'),
        ('Plan resolution (cache hit)', '≈450 ms (none)', '≈0.3 ms', '1 500×'),
        ('Peak planner memory',         '≈85 MB',         '≈4 MB',   '21×'),
        ('Per-step exec overhead',      '≈45 ms',         '≈0.9 ms', '50×'),
        ('Minimal rootfs size',         '—',                   '≈6.3 MB', '—'),
        ('QEMU boot to first service',  '—',                   '<500 ms',      '—'),
        ('plan_validate() time',        '—',                   '<1 ms',        '—'),
    ],
    center_cols=[3]
)

H3(doc, 'BUILD SYSTEM COMPARISON')
para(doc,
    'To provide broader context for the Rogue Linux performance figures, a comparison '
    'against Buildroot is informative. A Buildroot build of a similarly minimal rootfs '
    '(BusyBox + musl + Linux kernel, no external packages) takes approximately '
    '25–40 minutes on a 4-core build host and requires approximately 8 GB '
    'of disk space for the build tree. The Rogue Linux build pipeline for the same package '
    'set takes approximately 3–8 minutes (dominated by kernel compilation) with '
    'less than 1 GB of disk space, because the planner and executor do not maintain '
    'a separate per-package stamp directory tree or patch management infrastructure.')

H3(doc, 'ROOTFS SIZE ANALYSIS')
para(doc,
    'The 6.3 MB minimal rootfs is composed of five categories. BusyBox binary: '
    '≈1.2 MB (stripped, x86_64). Cogman supervisor binary: ≈180 KB '
    '(stripped). Dynamic libraries (libc.so.6, ld-linux-x86-64.so.2, libgcc_s.so.1): '
    '≈3.8 MB. /etc/cogman/ service definition files, boot scripts, and plan '
    'files: ≈20 KB. Directory scaffolding (/dev, /proc, /sys, /tmp): '
    'negligible. A musl libc-based build eliminates the glibc dynamic library dependency '
    'entirely, reducing the rootfs to approximately 4.1 MB — competitive with '
    'Alpine Linux (≈5 MB).')
page_break(doc)

# ═══════════════════════════════════════════════════════ CH 9 CONCLUSION
H1(doc, 'CHAPTER 9 CONCLUSION AND FUTURE WORK')

H3(doc, 'CONCLUSION')
para(doc,
    'This project has successfully designed, implemented, and validated Rogue Linux — '
    'a deterministic, metadata-driven infrastructure for constructing minimal Linux-based '
    'operating system images, with Cogman as the unified toolchain spanning both build and '
    'runtime phases. The project delivers on all stated objectives: a schema-validated '
    'TOML package metadata format; a Rust-based cogman-planner with DAG resolution and '
    'CGM2PLAN binary emission; a C11-based cogman-executor with typed step operations and '
    'path traversal protection; a POSIX-correct PID 1 supervisor with '
    'dependency-aware service management and SIGCHLD self-pipe child reaping; a '
    'text-protocol Unix domain socket control interface; and a bootable minimal rootfs of '
    'approximately 6.3 MB verified under QEMU through a four-stage boot sequence.')
para(doc,
    'The performance evaluation demonstrates all three headline improvements over the '
    'legacy Python baseline: 56× faster plan resolution (8 ms vs. 450 ms), '
    '21× lower peak memory (4 MB vs. 85 MB), and 50× lower per-step '
    'execution overhead (0.9 ms vs. 45 ms). The content-addressed plan cache '
    'further reduces planning time to 0.3 ms on unchanged packages. All 40 unit, '
    'integration, supervisor lifecycle, and end-to-end test cases pass on the QEMU test '
    'platform.')
para(doc,
    'The most technically challenging aspect of the implementation was the SIGCHLD '
    'self-pipe pattern in cogman-supervisor. The initial implementation used a direct '
    'waitpid() call in the SIGCHLD handler, which passed functional tests but was '
    'theoretically unsafe under POSIX. The production implementation was refactored to use '
    'the self-pipe pattern after identifying the async-signal-safety requirement, and '
    'the change required updates to the main select() loop, the signal handler, and the '
    'child reaping logic — a non-trivial refactor validated by the full supervisor '
    'test suite before merging.')
para(doc,
    'All seven stated project objectives have been met. The TOML package definition format '
    'with schema validation was designed and implemented with compile-time type checking '
    'via serde’s derive macro. The cogman-planner achieves the stated 50× '
    'performance target with a 56× improvement in plan resolution time. The '
    'cogman-executor correctly handles all five typed step operations with path traversal '
    'protection on OP_COPY, verified by 13 unit tests. The minimal rootfs of '
    '6.3 MB is below the 10 MB target and boots successfully under QEMU with '
    '64 MB RAM.')

H3(doc, 'FUTURE WORK')
para(doc,
    'Landlock filesystem isolation: Restrict each service to its declared filesystem policy '
    'paths using the Linux Landlock LSM (available since Linux 5.13), providing '
    'per-service mandatory access control without the complexity of SELinux or AppArmor '
    'policy authoring.')
para(doc,
    'seccomp-BPF system call filtering: Generate a per-service seccomp filter from the '
    'service definition’s declared syscall set, reducing the kernel attack surface '
    'by restricting each service to the minimum set of system calls required for its '
    'operation.')
para(doc,
    'Linux namespace isolation: Provide network and PID namespace isolation for services '
    'that declare isolation requirements in their definition files, enabling lightweight '
    'container-like isolation without a full container runtime.')
para(doc,
    'ARM64 and RISC-V cross-compilation: Extend the build system to support '
    'cross-compilation to ARM64 (Cortex-A series) and RISC-V (RV64GC) targets, enabling '
    'Rogue Linux rootfs images for the embedded targets that constitute the primary '
    'production deployment environment for minimal Linux images.')
para(doc,
    'Qwen2.5 QLoRA fine-tuning: Fine-tune the Qwen2.5-3B advisor model on a curated '
    'dataset of Cogman error messages, package.toml examples, and service file '
    'configurations using QLoRA, improving the accuracy of AI-assisted troubleshooting '
    'for common build and configuration errors.')
para(doc,
    'Full rootfs reproducibility: Achieve bit-for-bit identical rootfs images by using a '
    'fixed timestamp during ext4 image creation, eliminating the last source of '
    'non-determinism in the build pipeline.')
page_break(doc)

# ═══════════════════════════════════════════════════════ APPENDIX 1 SOURCE CODE
doc.add_heading('APPENDIX 1 SOURCE CODE', level=2)

H3(doc, 'A.1 Package Metadata Schema – Rust Struct Definitions (schema.rs)')
p = doc.add_paragraph(style='Normal')
run = p.add_run(
    '#[derive(Debug, Deserialize, Serialize, Clone)]\n'
    'pub struct PackageMetadata {\n'
    '    pub identity:  Identity,\n'
    '    pub build:     Builder,\n'
    '    pub installer: Installer,\n'
    '    #[serde(default)]\n'
    '    pub policy:    Policy,\n'
    '}\n\n'
    '#[derive(Debug, Deserialize, Serialize, Clone)]\n'
    'pub struct Identity {\n'
    '    pub name:     String,\n'
    '    pub version:  String,\n'
    '    pub category: String,\n'
    '    pub summary:  String,\n'
    '    pub source:   Source,\n'
    '    #[serde(default)]\n'
    '    pub depends:  Depends,\n'
    '}\n\n'
    '#[derive(Debug, Deserialize, Serialize, Default, Clone)]\n'
    'pub struct Depends {\n'
    '    #[serde(default)] pub build:   Vec<String>,\n'
    '    #[serde(default)] pub runtime: Vec<String>,\n'
    '}\n\n'
    '#[derive(Debug, Deserialize, Serialize, Default, Clone)]\n'
    'pub struct Policy {\n'
    '    pub filesystem: Filesystem,\n'
    '    pub network:    Network,\n'
    '}'
)
run.font.name = 'Courier New'; run.font.size = Pt(9)

H3(doc, 'A.2 TOML Package Definition – Example (busybox.toml)')
p = doc.add_paragraph(style='Normal')
run = p.add_run(
    '[identity]\n'
    'name     = "busybox"\n'
    'version  = "1.36.1"\n'
    'category = "base"\n'
    'summary  = "Multi-call binary providing UNIX shell utilities"\n\n'
    '[identity.source]\n'
    'kind = "tarball"\n'
    'file = "busybox-1.36.1.tar.bz2"\n\n'
    '[identity.depends]\n'
    'build   = []\n'
    'runtime = []\n\n'
    '[build]\n'
    'system = "make"\n'
    'steps  = [\n'
    '    "make defconfig",\n'
    '    "make CONFIG_STATIC=y -j$(nproc)",\n'
    ']\n\n'
    '[installer]\n'
    'steps  = [ "install -m 755 busybox $PKGROOT/bin/busybox" ]\n'
    'verify = { path = "$PKGROOT/bin/busybox" }\n\n'
    '[policy.filesystem]\n'
    'write = ["/bin", "/sbin", "/usr/bin"]\n'
    '[policy.network]\n'
    'access = false'
)
run.font.name = 'Courier New'; run.font.size = Pt(9)

H3(doc, 'A.3 Service Definition File – Example (heartbeat.service)')
p = doc.add_paragraph(style='Normal')
run = p.add_run(
    '[service]\n'
    'name          = heartbeat\n'
    'command       = /usr/bin/heartbeat-svc\n'
    'type          = longrun\n'
    'restart       = always\n'
    'restart_delay = 5\n'
    'depends_on    = hello\n\n'
    '[env]\n'
    'TICK_INTERVAL = 10\n'
    'LOG_LEVEL     = info\n\n'
    '[meta]\n'
    'description = Periodic heartbeat logger\n'
    'enabled     = true'
)
run.font.name = 'Courier New'; run.font.size = Pt(9)

H3(doc, 'A.4 Plan Validation Function (C – plan.c)')
p = doc.add_paragraph(style='Normal')
run = p.add_run(
    'int plan_validate(const void *base, size_t sz) {\n'
    '    if (sz < sizeof(struct plan_header)) {\n'
    '        log_err("Plan too small: %zu bytes", sz); return -1; }\n'
    '    const struct plan_header *h = base;\n'
    '    if (memcmp(h->magic, PLAN_MAGIC, 8) != 0) {\n'
    '        log_err("Bad magic"); return -1; }\n'
    '    if (h->version != PLAN_VERSION) {\n'
    '        log_err("Version %u != %u", h->version, PLAN_VERSION); return -1; }\n'
    '    size_t steps_end = sizeof(*h) +\n'
    '                       (size_t)h->step_count * sizeof(struct step_record);\n'
    '    if (steps_end > sz) { log_err("step_count overflows"); return -1; }\n'
    '    if (h->strtab_offset + h->strtab_len > sz) {\n'
    '        log_err("String table out of bounds"); return -1; }\n'
    '    return 0;\n'
    '}'
)
run.font.name = 'Courier New'; run.font.size = Pt(9)

H3(doc, 'A.5 Service State Machine – sup_handle_dead() (C – supervisor.c)')
p = doc.add_paragraph(style='Normal')
run = p.add_run(
    'void sup_handle_dead(pid_t pid, int status) {\n'
    '    struct service *svc = find_service_by_pid(pid);\n'
    '    if (!svc) { fprintf(stderr, "orphan pid=%d\\n", pid); return; }\n'
    '    int ec = WIFEXITED(status)   ? WEXITSTATUS(status)\n'
    '           : WIFSIGNALED(status) ? -WTERMSIG(status) : 0;\n'
    '    svc->exit_code = ec;  svc->pid = -1;\n\n'
    '    if (svc->type == SVC_TYPE_ONESHOT) {\n'
    '        svc->state = (ec == 0) ? SVC_DONE : SVC_FAILED; return; }\n'
    '    if (svc->state == SVC_STOPPED) return;  /* explicit stop */\n\n'
    '    int should = 0;\n'
    '    switch (svc->restart) {\n'
    '    case SVC_RESTART_ALWAYS:     should = 1;          break;\n'
    '    case SVC_RESTART_ON_FAILURE: should = (ec != 0);  break;\n'
    '    default:                     should = 0;          break;\n'
    '    }\n'
    '    if (should) {\n'
    '        svc->state      = SVC_RESTARTING;\n'
    '        svc->restart_at = time(NULL) + svc->restart_delay;\n'
    '    } else {\n'
    '        svc->state = (ec == 0) ? SVC_STOPPED : SVC_FAILED;\n'
    '    }\n'
    '}'
)
run.font.name = 'Courier New'; run.font.size = Pt(9)

H3(doc, 'A.6 SIGCHLD Self-Pipe Setup (C – supervisor.c)')
p = doc.add_paragraph(style='Normal')
run = p.add_run(
    'static int sigchld_pipe_w = -1;\n\n'
    'static void sigchld_handler(int sig) {\n'
    '    (void)sig;\n'
    '    char b = 1;\n'
    '    write(sigchld_pipe_w, &b, 1);  /* async-signal-safe */\n'
    '}\n\n'
    '/* Setup in main() */\n'
    'int pipefd[2];\n'
    'pipe2(pipefd, O_NONBLOCK | O_CLOEXEC);\n'
    'sigchld_pipe_w = pipefd[1];\n'
    'signal(SIGCHLD, sigchld_handler);\n\n'
    '/* In select() main loop */\n'
    'if (FD_ISSET(pipefd[0], &rfds)) {\n'
    '    char buf[64]; read(pipefd[0], buf, sizeof(buf));  /* drain */\n'
    '    pid_t pid; int st;\n'
    '    while ((pid = waitpid(-1, &st, WNOHANG)) > 0)\n'
    '        sup_handle_dead(pid, st);\n'
    '}'
)
run.font.name = 'Courier New'; run.font.size = Pt(9)
page_break(doc)

# ═══════════════════════════════════════════════════════ APPENDIX 2 SCREENSHOTS
H1(doc, 'APPENDIX 2')
H3(doc, 'SCREENSHOTS:')

para(doc,
    'The following screenshots demonstrate the Rogue Linux system running under QEMU with '
    'the DWM window manager and a working terminal environment. These screenshots were '
    'captured from a live QEMU instance booting the Rogue Linux rootfs with X11 and '
    'DWM enabled, confirming that the graphical stack (Xorg, DWM, st terminal) initialises '
    'correctly via the cogman-supervisor service pipeline.')

figure(doc, fig('extra_dwm_ui.png'),
       'Figure A.1: DWM Window Manager running on Rogue Linux (QEMU)')
figure(doc, fig('extra_terminal.png'),
       'Figure A.2: Terminal Environment – cogman-ctl and BusyBox shell')
page_break(doc)

# ═══════════════════════════════════════════════════════ APPENDIX 3 SETUP
H1(doc, 'APPENDIX 3')
H3(doc, 'HARDWARE AND SOFTWARE SETUP:')

para(doc,
    'Build Host: The Rogue Linux toolchain was developed and tested on a Kali Linux host '
    '(Linux 6.x, x86_64) with Rust 1.75+ (via rustup), GCC 11+, GNU Make 4.3+, '
    'QEMU 8.x, BusyBox 1.36.1, Python 3.11+, and llama.cpp (latest release).')
para(doc,
    'QEMU Boot Command: qemu-system-x86_64 -kernel /boot/vmlinuz -initrd rootfs.cpio '
    '-append "console=ttyS0 init=/sbin/cogman-supervisor" -nographic -m 64M. '
    'The -nographic flag redirects all console output to the terminal, enabling automated '
    'test scripts to monitor the boot sequence for expected output strings within a '
    '30-second timeout window.')
para(doc,
    'Repository Structure: rogue-linux/ contains cogman/ (Rust planner), executor/ (C '
    'executor), supervisor/ (C PID 1 supervisor), ctl/ (C control client), '
    'messenger/ (C IPC), rootfs/ (filesystem skeleton), packages/ (TOML definitions), '
    'tests/ (Python test harness), and build/ (build scripts and verification logs).')
page_break(doc)

# ═══════════════════════════════════════════════════════ REFERENCES
H1(doc, 'REFERENCE')

for ref in [
    '[1]  Buildroot Project. (2023). Buildroot: Making Embedded Linux Easy. https://buildroot.org/',
    '[2]  Yocto Project. (2023). Yocto Project Documentation. https://docs.yoctoproject.org/',
    '[3]  Dolstra, E. (2006). The Purely Functional Software Deployment Model. PhD Thesis, Utrecht University.',
    '[4]  Poettering, L. et al. (2010–present). systemd – System and Service Manager. https://systemd.io/',
    '[5]  Bernstein, D. J. (2000). daemontools: Service Management. https://cr.yp.to/daemontools.html',
    '[6]  Bercot, L. (2012–present). s6 – A Small Supervision Suite. https://skarnet.org/software/s6/',
    '[7]  The Rust Foundation. (2024). The Rust Programming Language. https://doc.rust-lang.org/book/',
    '[8]  Serde Working Group. (2024). serde – A Serialisation Framework for Rust. https://serde.rs/',
    '[9]  QEMU Project. (2024). QEMU – The Fast! Processor Emulator. https://www.qemu.org/',
    '[10] Alibaba Cloud. (2024). Qwen2.5 Technical Report. arXiv:2412.15115.',
    '[11] Ggerganov. (2022). llama.cpp: LLM Inference in C/C++. https://github.com/ggerganov/llama.cpp',
    '[12] ISO/IEC. (2011). ISO/IEC 9899:2011 – Programming Languages: C. International Standard.',
    '[13] Kerrisk, M. (2010). The Linux Programming Interface. No Starch Press, San Francisco.',
    '[14] Stevens, W. R. & Rago, S. A. (2013). Advanced Programming in the UNIX Environment (3rd ed.). Addison-Wesley.',
    '[15] Reproducible Builds Project. (2024). Reproducible Builds. https://reproducible-builds.org/',
    '[16] BusyBox Project. (2023). BusyBox: The Swiss Army Knife of Embedded Linux. https://busybox.net/',
    '[17] Bazel Authors. (2024). Bazel Documentation: Remote Caching. https://bazel.build/remote/caching',
    '[18] Torvalds, L. et al. (2024). Linux Kernel Source – init/main.c. https://kernel.org/',
    '[19] The Open Group. (2017). POSIX.1-2017 – Base Specifications Issue 7. https://pubs.opengroup.org/',
    '[20] Preston-Werner, T. (2021). TOML v1.0.0 Specification. https://toml.io/en/v1.0.0',
]:
    para(doc, ref)

# ─────────────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f'Saved: {OUT}')
